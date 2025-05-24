"""
Word document processor with memory-efficient streaming capabilities.
"""
import os
import io
import re
import time
import tempfile
from typing import Dict, Generator, List, Optional, Any, Union, Tuple

import pandas as pd
import numpy as np
from tqdm import tqdm
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

from app.utils.logging_config import get_logger
from app.processors.base_processor import BaseProcessor

logger = get_logger(__name__)

class DocxProcessor(BaseProcessor):
    """
    Process large Word (DOCX) files with memory optimization.
    
    This class handles Word documents efficiently by:
    1. Processing document in sections to minimize memory usage
    2. Extracting tables and text with streaming approach
    3. Converting document content to structured data
    """
    
    def __init__(self, **kwargs):
        """Initialize the DOCX processor with shared parameters."""
        super().__init__(**kwargs)
        # DOCX-specific settings
        self.extract_tables = kwargs.get('extract_tables', True)
        self.extract_text = kwargs.get('extract_text', True)
        self.extract_images = kwargs.get('extract_images', False)
        self.extract_headers = kwargs.get('extract_headers', True)
        self.temp_dir = kwargs.get('temp_dir', tempfile.gettempdir())
        self.max_section_size = kwargs.get('max_section_size', 50)  # Process document in sections
        
    def read_file(self, file_path: str) -> Generator[pd.DataFrame, None, None]:
        """
        Read a Word document in memory-efficient chunks.
        
        Args:
            file_path: Path to the DOCX file
            
        Yields:
            DataFrame chunks from the document
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
            
        # Get file info
        file_info = self.get_file_info(file_path)
        
        try:
            # Open the document
            doc = Document(file_path)
            
            # Extract tables if requested
            if self.extract_tables:
                tables = self._extract_tables(doc)
                if tables:
                    for i, table_df in enumerate(tables):
                        # Add metadata
                        table_df['table_num'] = i + 1
                        
                        # Optimize data types
                        try:
                            # Use base class method to detect optimal types
                            optimal_types = self.detect_data_type(table_df)
                            # Apply the types
                            for col, dtype in optimal_types.items():
                                if col in table_df.columns:
                                    try:
                                        table_df[col] = table_df[col].astype(dtype)
                                    except:
                                        pass  # Skip if conversion fails
                            optimized_df = table_df
                        except:
                            optimized_df = table_df  # Use original if optimization fails
                        
                        yield optimized_df
                        
                        # Clean up memory periodically
                        if (i + 1) % 5 == 0:
                            self.memory_monitor.cleanup_memory()
            
            # Extract text if requested
            if self.extract_text:
                # Process document in sections to minimize memory usage
                sections = []
                current_section = []
                
                # Process paragraphs
                for i, para in enumerate(doc.paragraphs):
                    # Add paragraph to current section
                    current_section.append({
                        'paragraph_num': i + 1,
                        'text': para.text,
                        'style': para.style.name if para.style else 'Normal',
                        'is_heading': para.style.name.startswith('Heading') if para.style else False
                    })
                    
                    # If we've reached max section size, process this section
                    if len(current_section) >= self.max_section_size:
                        sections.append(current_section)
                        current_section = []
                
                # Add the last section if it has content
                if current_section:
                    sections.append(current_section)
                
                # Process each section
                for i, section in enumerate(sections):
                    # Convert section to DataFrame
                    section_df = pd.DataFrame(section)
                    
                    # Add section metadata
                    section_df['section_num'] = i + 1
                    
                    # Yield this section
                    yield section_df
                    
                    # Clean up memory after each section
                    self.memory_monitor.cleanup_memory()
                    
        except Exception as e:
            logger.error(f"Error reading DOCX file {file_path}: {str(e)}")
            # Return error DataFrame
            error_df = pd.DataFrame({
                'error': [str(e)],
                'file_path': [file_path]
            })
            yield error_df
            
    def get_file_info(self, file_path: str) -> Dict[str, Any]:
        """
        Get metadata about a Word document without reading its full contents.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary with file metadata
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
            
        file_info = {
            "file_path": file_path,
            "file_name": os.path.basename(file_path),
            "file_size_bytes": os.path.getsize(file_path),
            "file_size_mb": os.path.getsize(file_path) / (1024 * 1024),
            "file_type": "docx",
        }
        
        try:
            # Open document to extract metadata
            doc = Document(file_path)
            
            # Count paragraphs and tables
            file_info["paragraph_count"] = len(doc.paragraphs)
            file_info["table_count"] = len(doc.tables)
            
            # Extract document properties if available
            core_properties = doc.core_properties
            file_info["metadata"] = {
                "author": core_properties.author or "",
                "title": core_properties.title or "",
                "subject": core_properties.subject or "",
                "created": str(core_properties.created) if core_properties.created else "",
                "modified": str(core_properties.modified) if core_properties.modified else "",
                "last_modified_by": core_properties.last_modified_by or ""
            }
            
            # Count headings by level
            heading_counts = {}
            for para in doc.paragraphs:
                style = para.style.name if para.style else "Normal"
                if style.startswith("Heading"):
                    heading_level = style.replace("Heading ", "")
                    heading_counts[heading_level] = heading_counts.get(heading_level, 0) + 1
            
            file_info["heading_counts"] = heading_counts
            
            # Get text sample from beginning of document
            text_sample = ""
            for i, para in enumerate(doc.paragraphs):
                if i < 5:  # Just take first 5 paragraphs
                    text_sample += para.text + "\n"
                else:
                    break
            
            file_info["text_sample"] = text_sample[:500] + "..." if len(text_sample) > 500 else text_sample
            
            # Analyze tables if any
            tables_info = []
            for i, table in enumerate(doc.tables):
                if i < 5:  # Just analyze first 5 tables
                    tables_info.append({
                        "table_num": i + 1,
                        "rows": len(table.rows),
                        "columns": len(table.columns) if table.rows else 0,
                    })
                    
            file_info["tables_info"] = tables_info
            
            # Estimate memory requirements
            file_info["estimated_memory_mb"] = self.estimate_memory_requirement(file_path)
            
            return file_info
            
        except Exception as e:
            logger.error(f"Error getting DOCX file info for {file_path}: {str(e)}")
            file_info["error"] = str(e)
            # Return basic file info even if detailed analysis fails
            return file_info
            
    def _extract_tables(self, doc: Document) -> List[pd.DataFrame]:
        """
        Extract tables from a Word document as DataFrames.
        
        Args:
            doc: The Word document
            
        Returns:
            List of DataFrames representing tables
        """
        tables_dfs = []
        
        for table_num, table in enumerate(doc.tables):
            try:
                # Get table data
                table_data = []
                
                # Process header row first if has_header
                headers = []
                if len(table.rows) > 0:
                    # Assuming first row is header
                    header_row = table.rows[0]
                    headers = [cell.text.strip() for cell in header_row.cells]
                
                # Process rows
                for row_idx, row in enumerate(table.rows):
                    # Skip header row since we already processed it
                    if row_idx == 0 and headers:
                        continue
                        
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                
                # Convert to DataFrame
                if table_data:
                    # Use headers if available, otherwise use default column names
                    if headers:
                        # Ensure headers have unique names
                        unique_headers = []
                        header_counts = {}
                        
                        for header in headers:
                            if header in header_counts:
                                header_counts[header] += 1
                                unique_headers.append(f"{header}_{header_counts[header]}")
                            else:
                                header_counts[header] = 0
                                unique_headers.append(header)
                                
                        # Create DataFrame with headers
                        df = pd.DataFrame(table_data, columns=unique_headers)
                    else:
                        # Create DataFrame without headers
                        df = pd.DataFrame(table_data)
                    
                    # Add table metadata
                    df['table_num'] = table_num + 1
                    
                    tables_dfs.append(df)
                    
            except Exception as e:
                logger.warning(f"Error extracting table {table_num}: {str(e)}")
                
        return tables_dfs
        
    def extract_document_structure(self, file_path: str) -> Dict[str, Any]:
        """
        Extract document structure with headings and sections.
        
        Args:
            file_path: Path to the DOCX file
            
        Returns:
            Dictionary with document structure
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"DOCX file not found: {file_path}")
            
        try:
            # Open the document
            doc = Document(file_path)
            
            # Track document structure
            structure = {
                "title": "",
                "headings": [],
                "sections": []
            }
            
            current_heading = None
            current_section = {
                "heading": "",
                "level": 0,
                "content": []
            }
            
            # Process each paragraph
            for para in doc.paragraphs:
                style = para.style.name if para.style else "Normal"
                text = para.text.strip()
                
                # Skip empty paragraphs
                if not text:
                    continue
                
                # Check if this is a heading
                if style.startswith("Heading") or style == "Title":
                    # Get heading level (Title is treated as level 0)
                    level = 0 if style == "Title" else int(style.replace("Heading ", ""))
                    
                    # If this is the title and we don't have one yet, set it
                    if style == "Title" and not structure["title"]:
                        structure["title"] = text
                    
                    # If we have a current section with content, save it
                    if current_section["content"]:
                        structure["sections"].append(current_section)
                    
                    # Start a new section
                    current_heading = {
                        "text": text,
                        "level": level
                    }
                    structure["headings"].append(current_heading)
                    
                    current_section = {
                        "heading": text,
                        "level": level,
                        "content": []
                    }
                else:
                    # Regular paragraph - add to current section
                    current_section["content"].append(text)
            
            # Add the last section if it has content
            if current_section["content"]:
                structure["sections"].append(current_section)
                
            return structure
            
        except Exception as e:
            logger.error(f"Error extracting document structure: {str(e)}")
            return {
                "error": str(e),
                "file_path": file_path
            }
            
    def extract_images(self, file_path: str, output_dir: Optional[str] = None) -> List[str]:
        """
        Extract images from Word document.
        
        Args:
            file_path: Path to the DOCX file
            output_dir: Directory to save images (None for temp dir)
            
        Returns:
            List of paths to extracted images
        """
        if output_dir is None:
            output_dir = os.path.join(self.temp_dir, "docx_images")
            
        os.makedirs(output_dir, exist_ok=True)
        
        image_paths = []
        
        try:
            # Since the python-docx library doesn't provide a direct way to extract images,
            # we'll use a workaround by treating the .docx as a zip file
            import zipfile
            from xml.etree.ElementTree import parse
            
            # Open the .docx file as a zip
            with zipfile.ZipFile(file_path) as docx_zip:
                # Get list of image files in the document
                image_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
                
                # Extract each image
                for img_path in image_files:
                    try:
                        # Get image filename
                        img_filename = os.path.basename(img_path)
                        
                        # Extract to output directory
                        output_path = os.path.join(output_dir, img_filename)
                        
                        # Extract the image
                        with open(output_path, 'wb') as img_file:
                            img_file.write(docx_zip.read(img_path))
                            
                        image_paths.append(output_path)
                        
                    except Exception as e:
                        logger.warning(f"Error extracting image {img_path}: {str(e)}")
                        
            logger.info(f"Extracted {len(image_paths)} images from {file_path}")
            return image_paths
            
        except Exception as e:
            logger.error(f"Error extracting images from DOCX: {str(e)}")
            return image_paths

    def process_file(self, file_path: str, **kwargs) -> Dict[str, Any]:
        """
        Process a DOCX file using the provided function or default processing.
        This method implements the abstract method from BaseProcessor.
        
        Args:
            file_path: Path to the DOCX file to process
            **kwargs: Additional processing parameters including:
                - process_fn: Function to apply to each chunk
                - output_path: Path to save processed output
                
        Returns:
            Dictionary with processing results
        """
        start_time = time.time()
        
        try:
            # Extract parameters
            process_fn = kwargs.get('process_fn')
            output_path = kwargs.get('output_path')
            
            # Initialize results
            total_rows_processed = 0
            processed_chunks = []
            
            # Track memory usage
            self.memory_monitor.start_tracking_step(f"docx_processing_{os.path.basename(file_path)}")
            
            # Process file in chunks
            for chunk in self.read_file(file_path):
                chunk_rows = len(chunk)
                
                # Apply processing function if provided
                if process_fn:
                    try:
                        chunk = process_fn(chunk)
                    except Exception as e:
                        logger.warning(f"Error applying process function to chunk: {str(e)}")
                
                processed_chunks.append(chunk)
                total_rows_processed += chunk_rows
                
                # Update memory peak tracking
                self.memory_monitor.update_step_peak(f"docx_processing_{os.path.basename(file_path)}")
            
            # Combine results if we have chunks to process
            if processed_chunks:
                combined_df = pd.concat(processed_chunks, ignore_index=True)
                
                # Save output if path provided
                if output_path:
                    combined_df.to_csv(output_path, index=False)
                    logger.info(f"Saved processed DOCX to {output_path}")
            else:
                combined_df = pd.DataFrame()
            
            # End memory tracking
            memory_stats = self.memory_monitor.end_tracking_step(f"docx_processing_{os.path.basename(file_path)}")
            
            # Calculate processing time
            processing_time = time.time() - start_time
            
            return {
                "status": "completed",
                "file_path": file_path,
                "output_path": output_path,
                "data": combined_df,
                "stats": {
                    "total_rows_processed": total_rows_processed,
                    "total_chunks": len(processed_chunks),
                    "processing_time": processing_time,
                    "memory_peak": memory_stats.get("peak_mb", 0),
                    "file_size_mb": os.path.getsize(file_path) / (1024 * 1024)
                }
            }
            
        except Exception as e:
            logger.error(f"Error processing DOCX file {file_path}: {str(e)}")
            return {
                "status": "failed",
                "error": str(e),
                "file_path": file_path,
                "stats": {
                    "total_rows_processed": 0,
                    "processing_time": time.time() - start_time,
                    "memory_peak": 0
                }
            }